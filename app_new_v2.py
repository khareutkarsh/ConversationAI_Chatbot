import os
import pickle
import streamlit as st
import torch
import numpy as np
import re
from transformers import pipeline
from rank_bm25 import BM25Okapi
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from docx import Document as DocxDocument

# ---------------------- STEP 1: SECURITY GUARDRAILS ---------------------- #
BLOCKED_TOPICS = ["violence", "hate speech", "politics", "illegal", "criminal activity"]

def is_valid_query(query):
    """Check if the query contains prohibited content."""
    return not any(blocked in query.lower() for blocked in BLOCKED_TOPICS)

# ---------------------- STEP 2: DATA LOADING ---------------------- #
def load_docx_from_folder(folder_path):
    """Load and extract text from all DOCX files in a folder."""
    all_text = []
    if not os.path.exists(folder_path):
        return "⚠️ Error: Document folder not found!"
    
    files = [f for f in os.listdir(folder_path) if f.endswith(".docx")]
    if not files:
        return "⚠️ Error: No financial reports found!"
    
    for file in files:
        file_path = os.path.join(folder_path, file)
        print(f"Processing: {file_path}")
        doc = DocxDocument(file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        all_text.append(text)
    
    combined_text = " ".join(all_text).strip()
    return combined_text if combined_text else "⚠️ Error: No valid text extracted from reports!"

docx_folder = "financial_statements"
financial_text = load_docx_from_folder(docx_folder)

if "⚠️" in financial_text:
    st.error(financial_text)
    st.stop()

# ---------------------- STEP 3: TEXT CHUNKING ---------------------- #
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
chunks = text_splitter.split_text(financial_text)
if not chunks:
    st.error("⚠️ Error: No chunks generated from financial documents!")
    st.stop()

print(f"✅ Total Chunks Processed: {len(chunks)}")

# ---------------------- STEP 4: VECTOR STORAGE ---------------------- #
embedding_model = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
BM25_PATH = "bm25_index.pkl"
CHROMA_DB_PATH = "chroma_db"

if os.path.exists(BM25_PATH) and os.path.exists(CHROMA_DB_PATH):
    print("✅ Loading saved models...")
    with open(BM25_PATH, "rb") as f:
        bm25 = pickle.load(f)
    vector_db = Chroma(persist_directory=CHROMA_DB_PATH, embedding_function=embedding_model)
else:
    print("⚠️ Processing and saving models for the first time...")
    documents = [Document(page_content=chunk) for chunk in chunks]
    vector_db = Chroma.from_documents(documents, embedding_model, persist_directory=CHROMA_DB_PATH)
    tokenized_chunks = [chunk.split() for chunk in chunks]
    bm25 = BM25Okapi(tokenized_chunks)
    with open(BM25_PATH, "wb") as f:
        pickle.dump(bm25, f)

# ---------------------- STEP 5: RETRIEVAL WITH GUARDRAILS ---------------------- #
def retrieve_bm25(query, top_k=3):
    """Retrieve financial text snippets using BM25 ranking."""
    if not is_valid_query(query):
        return ["⚠️ This query is blocked due to sensitive content."]
    
    tokenized_query = query.split()
    scores = bm25.get_scores(tokenized_query)
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_k]
    
    return [chunks[i] for i in top_indices] if top_indices else ["⚠️ No relevant financial data found."]

def hybrid_retrieve(query, k=3):
    """Retrieve financial data using BM25 and embeddings."""
    bm25_results = retrieve_bm25(query, top_k=k)
    if "⚠️" in bm25_results[0]:
        return bm25_results
    
    embedding_results = vector_db.similarity_search(query, k=k)
    return bm25_results + [doc.page_content for doc in embedding_results]

# ---------------------- STEP 6: MODEL GENERATION WITH FACT-CHECKING ---------------------- #
generator = pipeline("question-answering", model="deepset/roberta-base-squad2")

def extract_numbers(text):
    """Extract all numerical values from text."""
    return re.findall(r"\$?\d+(?:\.\d+)?", text)

def generate_response(query, context):
    """Generate a verified financial response based on retrieved data."""
    response = generator(question=query, context=context)
    return response['answer'].strip() if response['score'] > 0.5 else "I could not find relevant financial data."

def validate_response(response, retrieved_context):
    """Check if AI-generated numbers match retrieved document numbers."""
    context_numbers = set(extract_numbers(" ".join(retrieved_context)))
    response_numbers = set(extract_numbers(response))
    if response_numbers - context_numbers:
        return "⚠️ The AI-generated response contained incorrect numbers. Please verify manually."
    return response

# ---------------------- STEP 7: STREAMLIT UI ---------------------- #
st.title("📊 Financial RAG Chatbot")
st.write("Ask questions based on the company's financial statements.")

user_query = st.text_input("Enter your financial question:")

if st.button("Get Answer"):
    if not user_query.strip():
        st.error("⚠️ Please enter a valid financial question.")
    elif not is_valid_query(user_query):
        st.error("⚠️ This query is blocked due to sensitive content.")
    else:
        retrieved_context = hybrid_retrieve(user_query)
        if "⚠️" in retrieved_context[0]:
            st.error(retrieved_context[0])
        else:
            raw_answer = generate_response(user_query, " ".join(retrieved_context))
            final_answer = validate_response(raw_answer, retrieved_context)
            st.write("### ✅ Final Answer:")
            st.write(final_answer)
